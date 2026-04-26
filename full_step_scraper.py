#!/usr/bin/env python3
"""
Full Step-by-Step ImmoCalcul Scraper (Robust Version)
- Uses a persistent browser context to defeat bot detection.
- Employs human-like typing to trigger autocomplete suggestions.
- Defaults to headed mode for stability (use --headless to override).
- Retains all original features for capturing data, screenshots, and generating reports.

Dependencies:
  pip install playwright Pillow google-api-python-client google-auth-oauthlib python-docx img2pdf python-dotenv

First-time Playwright:
  python -m playwright install chromium
  
"""

import os
import re
import json
import argparse
import asyncio
import random
import sys
import traceback
import logging
import subprocess
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Any
import pikepdf
import tempfile
import shutil

import requests
from playwright.async_api import async_playwright, Page, TimeoutError as PWTimeoutError, Locator
from PIL import Image, ImageDraw, ImageFont
import img2pdf
from docx import Document

# --- Import and setup logging (expected to exist in your project) ---
from logger_config import set_step
from logger_config import setup_logging  # used in __main__

# --- Load .env file ---
from dotenv import load_dotenv
load_dotenv()

# --- Google Drive API imports for OAuth 2.0 ---
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.errors import HttpError

# If modifying these scopes, delete the file token.json.
SCOPES = ['https://www.googleapis.com/auth/drive']
TOKEN_PATH = Path('token.json')         # created after first authorization
CREDENTIALS_PATH = Path('credentials.json')
PARENT_FOLDER_ID = False #os.getenv("PARENT_DRIVE_FOLDER_ID")  # set via env or --parent-folder-id

CURRENCY_PATTERN = re.compile(r'(\d[\d\s\u00A0]*\d)\s*\$')
MAP_ITEMS = [
    "Zonage municipal",
    "Plan cadastral",
    "Aires protégées",
    "Zone inondable",
    "Milieu humide",
    "Glissement de terrain",
    "Zone Agricole",
    "Feux de forêt",
    "Terrains Contaminés",
]

# Risk analysis mapping
RISK_MAP = {
    "Présence zone inondable": "Zone innondable",
    "Zone exposée aux glissements de terrain": "Glissement terrain",
    "Présence de milieu humide": "Zone humide",
    "Présence zone agricole": "Zone agricole",
    "Contamination du terrain": "Terrain contaminé",
    "Présence d'aires Protégée": "Présence d'air protégé",
    "Feux de forêt": "Feu foret"
}

# Active tab content container
ACTIVE_SECTION_XPATH = "//*[contains(@class,'dataAdress_modalProperty') and contains(@class,'dataAdress_active')]"

# The inner block we need to capture completely
DATA_SECTION_XPATH = (
    "//main[@id='main']"
    "//div[contains(@class,'map_blocvueSession')]"
    "/div[contains(@class,'dataAdress_modalProperty') and contains(@class,'dataAdress_active')]"
    "/div[contains(@class,'dataAdress_outerContainer')]"
    "/div[contains(@class,'dataAdress_blocDataAdress')]"
)

LOT_SEARCH_PATH = "//div[contains(@class, 'map_rightSearch')]//span[contains(@class, 'map_TextListSearch')]//span[contains(@class, 'map_adresseproperty')]"
RESUME_XPATH = "//div[contains(@class,'dataAdress_tabContainer')]//div[2]//p[contains(@class,'dataAdress_resume')]"
CHAR6_XPATH = "(//div[contains(@class,'dataAdress_proprieteCaracteristiques')]//div[contains(@class,'dataAdress_caracteristiqueItem')])[6]//span[contains(@class,'dataAdress_caracteristiqueSubTitre')]"
MAIN_PHOTO_XPATH = "//div[@class='tablet computer']/img"
AUTOCOMPLETE_ADDR_XPATH = "//div[contains(@class,'ElasticSearchAutocomplete_resultItem')]"
RISK_ISSUES_XPATH = "//div[contains(@class,'dataAdress_iconAlertError')]/parent::div/following-sibling::div//h3"

# Comparables dialog selectors
COMPARABLES_TRIGGER_XPATH = "//main[@id='main']//div[contains(@class,'chart_oneButton')][1]"
COMPARABLES_DIALOG_CONTAINER_XPATH = "//div[contains(@class,'MuiDialog-container') and contains(@class,'MuiDialog-scrollPaper')]"
# exact absolute XPath (ensure starts with // so Playwright treats it as XPath)
COMPARABLES_DIALOG_CONTENT_XPATH_EXACT = (
    "//html/body"
    "/div[contains(@class,'MuiDialog-root')]"
    "/div[contains(@class,'MuiDialog-container') and contains(@class,'MuiDialog-scrollPaper')]"
    "/div[contains(@class,'MuiDialog-paper') or contains(@class,'MuiPaper-root')]"
    "/div[contains(@class,'MuiDialogContent-root')]"
)
# robust fallbacks
COMPARABLES_DIALOG_CONTENT_XPATH = (
    "//body/div[contains(@class,'MuiDialog-root')]"
    "//div[contains(@class,'MuiDialogContent-root')]"
)
COMPARABLES_DIALOG_PAPER_XPATH = (
    "//body/div[contains(@class,'MuiDialog-root')]"
    "//div[contains(@class,'MuiDialog-paper') or contains(@class,'MuiPaper-root')]"
)

SATELLITE_BTN_XPATH = "//div[contains(@class,'map_customMapControls')]//div[contains(@class,'map_customSatelliteBtn')][1]//span"
MAP_LAYER_BASE_XPATH = "((//div[contains(@class,'map_containerSwitchMapLayer')])[2]//div[contains(@class,'map_blocLayerImage')])[{}]"
MAP_CANVAS_XPATH = "//div[contains(@class,'ol-layer')]//canvas"


def get_drive_service():
    creds = None
    if TOKEN_PATH.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_PATH), SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not CREDENTIALS_PATH.exists():
                set_step("error")
                logging.error(f"Missing Google OAuth credentials file: {CREDENTIALS_PATH}")
                logging.error("Please download it from Google Cloud Console and place it in the script's directory.")
                sys.exit(1)
            flow = InstalledAppFlow.from_client_secrets_file(str(CREDENTIALS_PATH), SCOPES)
            creds = flow.run_local_server(port=0)
        with open(TOKEN_PATH, 'w') as token:
            token.write(creds.to_json())
    try:
        service = build('drive', 'v3', credentials=creds)
        return service
    except HttpError as error:
        set_step("error")
        logging.error(f'An error occurred with Google Drive service: {error}', exc_info=True)
        return None

def create_drive_folder(service, name, parent_id):
    file_metadata = {'name': name, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [parent_id]}
    folder = service.files().create(body=file_metadata, fields='id').execute()
    return folder.get('id')

def upload_file_to_drive(service, file_path, folder_id):
    if not Path(file_path).exists():
        logging.warning(f"File not found for upload: {file_path}")
        return None, None
    file_metadata = {'name': os.path.basename(file_path), 'parents': [folder_id]}
    media = MediaFileUpload(str(file_path), resumable=True)
    try:
        file = service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        file_id = file.get('id')
        # Make public
        service.permissions().create(fileId=file_id, body={'role': 'reader', 'type': 'anyone'}).execute()
        logging.info(f"Uploaded '{os.path.basename(file_path)}' to Drive. File ID: {file_id}")
        return file_id, file.get('webViewLink')
    except HttpError as error:
        set_step("drive-permission")
        logging.warning(f"Failed to upload {file_path}: {error}")
        return None, None

def get_folder_url(folder_id):
    return f"https://drive.google.com/drive/folders/{folder_id}"

def rand_delay(args):
    if args.fixed_delay is not None:
        return args.fixed_delay
    return random.uniform(args.delay_min, args.delay_max)

async def human_wait(args, label=""):
    await asyncio.sleep(rand_delay(args))


class StepCapture:
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.counter = 1

    async def _scroll_page_fully(self, page: Page):
        """Scroll just the window to trigger lazy content elsewhere."""
        try:
            total_height = await page.evaluate("document.body.scrollHeight")
            viewport_height = page.viewport_size['height']
            current_position = 0
            while current_position < total_height:
                await page.mouse.wheel(0, viewport_height)
                await page.wait_for_timeout(300)
                new_height = await page.evaluate("document.body.scrollHeight")
                if new_height == total_height:
                    break
                total_height = new_height
                current_position += viewport_height
            await page.evaluate("window.scrollTo(0, 0)")
            await page.wait_for_timeout(150)
        except Exception as e:
            logging.warning(f"Could not scroll page fully: {e}")

    async def shot(self, page: Page, tag: str, full_page=True):
        """Generic page-level screenshot; kept for optional use."""
        name = f"{self.counter:02d}_{tag}.png"
        path = self.base_dir / name
        try:
            if full_page:
                await self._scroll_page_fully(page)
            await page.screenshot(path=str(path), full_page=full_page)
            logging.info(f"Screenshot captured: {name}")
        except Exception as e:
            set_step("capture-failed")
            logging.warning(f"Screenshot failed ({tag}): {e}")
        self.counter += 1
        return path

    async def shot_xpath_zoomfit(self, page: Page, xpath: str, tag: str, zoom: float = 0.33, hide_fixed: bool = True):
        """
        Capture a tall element fully by shrinking the scroll container's content.
        - No Playwright scroll_into_view_if_needed (avoids 60s timeouts).
        - We move the page and container via JS (single adjustment), then do one element screenshot.
        """
        name = f"{self.counter:02d}_{tag}.png"
        out_path = self.base_dir / name
        locator = page.locator(xpath).first
        el = None
        try:
            # Only require 'attached' to avoid waiting on visibility that can be affected by CSS transforms
            await locator.wait_for(state="attached", timeout=30000)
            el = await locator.element_handle()
            if not el:
                raise RuntimeError("Element handle not found for zoomfit capture.")

            # Prepare zoom wrapper around the scroller's children and optionally hide fixed overlays
            prepare_ok = await el.evaluate(
                """(node, ctx) => {
                  const z = ctx.z;
                  const hideFixed = ctx.hideFixed;

                  const getScroller = (n) => {
                    let cur = n;
                    while (cur && cur !== document.body) {
                      const cs = getComputedStyle(cur);
                      const oy = cs.overflowY;
                      if ((oy === 'auto' || oy === 'scroll' || oy === 'overlay') && cur.scrollHeight > cur.clientHeight + 1) return cur;
                      cur = cur.parentElement;
                    }
                    return document.scrollingElement || document.documentElement;
                  };

                  const scroller = getScroller(node);
                  if (!scroller) return false;

                  const mark = (elem) => {
                    if (!elem.hasAttribute('data-prev-style')) {
                      elem.setAttribute('data-prev-style', elem.getAttribute('style') || '');
                    }
                  };

                  mark(scroller);

                  // Create wrapper and move all children into it
                  const wrapper = document.createElement('div');
                  wrapper.setAttribute('data-zoom-wrapper', '1');
                  wrapper.style.transformOrigin = 'top left';
                  wrapper.style.transform = `scale(${z})`;

                  const scRect = scroller.getBoundingClientRect();
                  const scWidth = scRect.width || scroller.clientWidth;
                  wrapper.style.width = (scWidth / z) + 'px';

                  const kids = Array.from(scroller.childNodes);
                  for (const ch of kids) wrapper.appendChild(ch);
                  scroller.appendChild(wrapper);

                  // Ensure scroller scrolls and starts at top
                  scroller.style.overflowY = 'auto';
                  scroller.scrollTop = 0;

                  // Hide fixed/sticky overlays outside scroller
                  if (hideFixed) {
                    const scrollerAncestors = new Set();
                    let p = scroller;
                    while (p) { scrollerAncestors.add(p); p = p.parentElement; }

                    const all = document.querySelectorAll('*');
                    for (const e of all) {
                      const pos = getComputedStyle(e).position;
                      if ((pos === 'fixed' || pos === 'sticky') && !scroller.contains(e) && !scrollerAncestors.has(e)) {
                        if (!e.hasAttribute('data-prev-style')) {
                          e.setAttribute('data-prev-style', e.getAttribute('style') || '');
                        }
                        e.style.display = 'none';
                        e.setAttribute('data-zoom-hidden', '1');
                      }
                    }
                  }

                  // Bring scroller itself near top of the page (single adjustment, not iterative)
                  const scRect2 = scroller.getBoundingClientRect();
                  window.scrollBy(0, scRect2.top - 80); // 80px pad for sticky header

                  // Align element near top inside the scroller (single adjustment)
                  let y = 0, cur = node;
                  while (cur && cur !== scroller) { y += cur.offsetTop || 0; cur = cur.offsetParent; }
                  scroller.scrollTop = Math.max(0, y - 20); // small pad

                  return true;
                }""",
                {"z": float(zoom), "hideFixed": bool(hide_fixed)}
            )
            if not prepare_ok:
                raise RuntimeError("Failed to prepare zoom wrapper.")

            # Give layout a moment; then capture the element directly without scrollIntoView
            await page.wait_for_timeout(180)

            # Use the element handle to avoid re-querying/visibility checks
            await el.screenshot(path=str(out_path))
            logging.info(f"Zoom-fit element screenshot saved: {out_path.name}")

        except Exception as e:
            set_step("capture-failed")
            logging.warning(f"Zoom-fit element screenshot failed ({tag}): {e}", exc_info=True)
        finally:
            # Restore DOM
            try:
                if el:
                    await el.evaluate(
                        """(node) => {
                          const getScroller = (n) => {
                            let cur = n;
                            while (cur && cur !== document.body) {
                              const cs = getComputedStyle(cur);
                              const oy = cs.overflowY;
                              if ((oy === 'auto' || oy === 'scroll' || oy === 'overlay') && cur.scrollHeight > cur.clientHeight + 1) return cur;
                              cur = cur.parentElement;
                            }
                            return document.scrollingElement || document.documentElement;
                          };
                          const scroller = getScroller(node);
                          if (!scroller) return;

                          const wrapper = scroller.querySelector('[data-zoom-wrapper=\"1\"]');
                          if (wrapper) {
                            const kids = Array.from(wrapper.childNodes);
                            for (const ch of kids) scroller.insertBefore(ch, wrapper);
                            wrapper.remove();
                          }

                          // Restore styles on all elements we touched
                          const modified = document.querySelectorAll('[data-prev-style], [data-zoom-hidden]');
                          for (const elem of modified) {
                            if (elem.hasAttribute('data-prev-style')) {
                              const prev = elem.getAttribute('data-prev-style') || '';
                              if (prev) elem.setAttribute('style', prev);
                              else elem.removeAttribute('style');
                              elem.removeAttribute('data-prev-style');
                            }
                            if (elem.hasAttribute('data-zoom-hidden')) {
                              elem.removeAttribute('data-zoom-hidden');
                            }
                          }
                        }"""
                    )
            except Exception:
                pass
            self.counter += 1

        return out_path


def parse_amounts(text: str) -> List[int]:
    if not text:
        return []
    cleaned = text.replace("\u00A0", " ")
    amounts = []
    for m in CURRENCY_PATTERN.findall(cleaned):
        try:
            amounts.append(int(re.sub(r"\s+", "", m)))
        except ValueError:
            pass
    return amounts

def classify(amounts: List[int], char6: str) -> str:
    if not amounts:
        return "Unknown"
    if len(amounts) == 1:
        v = amounts[0]
        return "C-Land" if v >= 1000000 else "R-Land"
    return "R-House"

def normalize_risk_heading(heading: str) -> str:
    """
    Strip leading/trailing spaces and periods, then map to canonical risk heading.
    If not found, return the cleaned heading.
    """
    h = heading.strip(" .")
    return RISK_MAP.get(h, h)

async def get_text_from_locator(locator: Optional[Locator]) -> str:
    if locator:
        try:
            return (await locator.inner_text()).strip()
        except Exception:
            pass
    return ""

async def get_all_text_active(page: Page) -> str:
    return await get_text_from_locator(page.locator(ACTIVE_SECTION_XPATH).first)

async def get_risk_issue_headings(page: Page) -> List[str]:
    """Get risk issue headings with normalized mapping."""
    out = []
    try:
        locators = await page.locator(RISK_ISSUES_XPATH).all()
        for l in locators:
            t = await get_text_from_locator(l)
            if t:
                out.append(normalize_risk_heading(t))
    except Exception:
        pass
    return out

async def capture_main_photo(page: Page, out_dir: Path, sc: StepCapture) -> Optional[str]:
    set_step("photo-main")
    try:
        img_el = page.locator(MAIN_PHOTO_XPATH).first
        if not await img_el.is_visible():
            logging.info("Main photo element not found.")
            return None
        raw_path = out_dir / "main_photo_raw.png"
        await img_el.screenshot(path=str(raw_path))
        final_path = out_dir / "main_photo.jpg"
        Image.open(raw_path).convert("RGB").save(final_path, quality=92)
        logging.info(f"Main photo captured and saved to {final_path}")
        return str(final_path)
    except Exception as e:
        logging.warning(f"Main photo capture failed: {e}")
        return None

def overlay_label(img_path: Path, label: str):
    """
    Larger, readable overlay label:
      - Dynamic font size based on image width (~2.8% of width; min 22, max 48)
      - Padding and rounded rectangle background
    """
    try:
        img = Image.open(img_path).convert("RGBA")
        draw = ImageDraw.Draw(img)

        w, h = img.size
        base_size = max(22, min(48, int(w * 0.028)))
        # Boldish font options fallback
        font = None
        for fname in ["arialbd.ttf", "Arial Bold.ttf", "DejaVuSans-Bold.ttf", "arial.ttf", "DejaVuSans.ttf"]:
            try:
                font = ImageFont.truetype(fname, base_size)
                break
            except IOError:
                continue
        if font is None:
            font = ImageFont.load_default()

        pad_x = max(12, base_size // 2)
        pad_y = max(8, base_size // 3)

        bbox = draw.textbbox((0, 0), label, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]

        rect_w = tw + pad_x * 2
        rect_h = th + pad_y * 2
        rect = [8, 8, 8 + rect_w, 8 + rect_h]

        try:
            draw.rounded_rectangle(rect, radius=max(8, base_size // 2), fill=(0, 0, 0, 176))
        except Exception:
            draw.rectangle(rect, fill=(0, 0, 0, 176))

        draw.text((8 + pad_x, 8 + pad_y), label, fill=(255, 255, 255, 255), font=font)

        img.save(img_path)
    except Exception as e:
        logging.warning(f"Label overlay failed for {img_path.name}: {e}")

## PDF Optimization settings (at the top with other constants)
PDF_TARGET_DPI = int(os.getenv("PDF_TARGET_DPI", "80"))  # 120 DPI for compressed images
PDF_JPEG_QUALITY = int(os.getenv("PDF_JPEG_QUALITY", "70"))  # 80 quality for compressed images

# Files that should NOT be compressed (keep original quality)
NO_COMPRESS_PATTERNS = [
    "01_main_tab_zoom33.png",
    "02_map_opened_zoom33.png",
    "03_avis_tab_zoom33.png", 
    "04_map_opened_zoom33.png"
]

def write_pdf_from_screenshots(out_dir: Path, summary: dict) -> Optional[Path]:
    """
    Creates optimized PDF from screenshots with selective compression.
    - Skips compression for main tabs (01, 03, 04) - keeps original quality
    - Compresses other images to 120 DPI with JPEG quality 80
    Uses pikepdf for final PDF compression.
    """
    set_step("maps-pdf")
    all_imgs: List[Path] = []
    seen = set()

    # Collect numbered screenshots
    for i in range(1, 300):
        for p in sorted(out_dir.glob(f"{i:02d}_*.png")):
            if p not in seen:
                all_imgs.append(p)
                seen.add(p)

    # Collect map layer screenshots
    for layer in summary.get("map_layers", []):
        p = Path(layer.get("file", ""))
        if p.exists() and p not in seen:
            all_imgs.append(p)
            seen.add(p)

    if not all_imgs:
        logging.info("No screenshots found to build PDF.")
        return None

    pdf_path = out_dir / "immocalcul.pdf"
    temp_dir = None
    optimized_images = []
    temp_pdf = None

    try:
        # Create temporary directory for optimized images
        temp_dir = Path(tempfile.mkdtemp(prefix="pdfopt_", dir=str(out_dir)))
        
        logging.info(f"Processing {len(all_imgs)} images for PDF creation...")
        
        # Optimize each image (selectively)
        for idx, img_path in enumerate(all_imgs, start=1):
            try:
                img = Image.open(img_path)
                filename = img_path.name
                
                # Check if this file should skip compression
                should_skip_compression = any(pattern in filename for pattern in NO_COMPRESS_PATTERNS)
                
                if should_skip_compression:
                    # Keep original quality - just convert to RGB and save as JPEG
                    logging.info(f"   Preserving original quality for: {filename}")
                    
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if img.mode in ('RGBA', 'LA'):
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img)
                        img = background
                    else:
                        img = img.convert('RGB')
                    
                    # Save at original DPI with maximum quality
                    temp_img_path = temp_dir / f"original_{idx:04d}.jpg"
                    original_dpi = img.info.get('dpi', (300, 300))
                    img.save(
                        temp_img_path,
                        'JPEG',
                        quality=95,  # High quality for non-compressed images
                        optimize=False,
                        dpi=original_dpi
                    )
                    optimized_images.append(str(temp_img_path))
                    
                else:
                    
                    # Convert to RGB (handle transparency)
                    if img.mode in ('RGBA', 'LA', 'P'):
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        if img.mode in ('RGBA', 'LA'):
                            background.paste(img, mask=img.split()[-1])
                        else:
                            background.paste(img)
                        img = background
                    else:
                        img = img.convert('RGB')
                    
                    # Scale to 120 DPI if current DPI is higher
                    current_dpi = img.info.get('dpi', (72, 72))[0]
                    
                    if current_dpi > PDF_TARGET_DPI:
                        scale = PDF_TARGET_DPI / current_dpi
                        new_size = (int(img.width * scale), int(img.height * scale))
                        img = img.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # Save compressed JPEG at 120 DPI with quality 80
                    temp_img_path = temp_dir / f"compressed_{idx:04d}.jpg"
                    img.save(
                        temp_img_path,
                        'JPEG',
                        quality=PDF_JPEG_QUALITY,
                        optimize=True,
                        progressive=True,
                        dpi=(PDF_TARGET_DPI, PDF_TARGET_DPI)
                    )
                    optimized_images.append(str(temp_img_path))
                
            except Exception as e:
                logging.warning(f"Failed to process {img_path.name}: {e}")
                # Fall back to original if processing fails
                optimized_images.append(str(img_path))
        
        # Create initial PDF from optimized images
        temp_pdf = out_dir / "temp_immocalcul.pdf"
        logging.info(f"Creating PDF from {len(optimized_images)} processed images...")
        
        with open(temp_pdf, 'wb') as f:
            f.write(img2pdf.convert(optimized_images))
        
        temp_size_mb = temp_pdf.stat().st_size / (1024 * 1024)
        logging.info(f"Initial PDF created: {temp_size_mb:.2f} MB")
        
        # Apply additional PDF compression with pikepdf
        logging.info("Applying PDF compression...")
        with pikepdf.open(temp_pdf) as pdf:
            pdf.save(
                pdf_path,
                compress_streams=True,
                stream_decode_level=pikepdf.StreamDecodeLevel.generalized,
                object_stream_mode=pikepdf.ObjectStreamMode.generate,
                linearize=True  # Optimize for web viewing
            )
        
        final_size_mb = pdf_path.stat().st_size / (1024 * 1024)
        logging.info(f"Optimized PDF created: {pdf_path} ({final_size_mb:.2f} MB)")
        
        if temp_size_mb > 0:
            reduction = ((temp_size_mb - final_size_mb) / temp_size_mb * 100)
            logging.info(f"Size reduction: {reduction:.1f}%")
        
        return pdf_path
        
    except Exception as e:
        set_step("pdf-merge-error")
        logging.warning(f"PDF creation failed: {e}", exc_info=True)
        return None
        
    finally:
        # Cleanup temporary files
        if temp_pdf and temp_pdf.exists():
            try:
                temp_pdf.unlink()
            except Exception:
                pass
        
        if temp_dir and temp_dir.exists():
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception:
                pass

def create_docx_summary(out_dir: Path, summary: dict) -> Optional[Path]:
    """Creates a .docx file with extracted text."""
    set_step("doc-export")
    doc_path = out_dir / "immocalcul.docx"
    try:
        doc = Document()
        doc.add_heading('ImmoCalcul Summary', 0)

        doc.add_heading('Property Details', level=1)
        doc.add_paragraph(f"Lot: {summary.get('lot', 'N/A')}")
        full_address = f"{summary.get('address_line_1','')} {summary.get('address_line_2','')}".strip()
        doc.add_paragraph(f"Address: {full_address if full_address else 'N/A'}")
        doc.add_paragraph(f"Classification: {summary.get('classification', 'N/A')}")

        sections = [
            ("Résumé (Resume)", summary.get("resume_text")),
            ("Main Active Section", summary.get("active_text_main")),
            ("Measures Section", summary.get("active_text_measures")),
            ("Avis (Notices) Section", summary.get("active_text_avis")),
            ("Risk Issues", "\n".join(summary.get("risk_issues", [])) if summary.get("risk_issues") else ""),
            ("Comparables", summary.get("comparables_text")),
        ]
        for title, content in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content if content else "N/A")

        doc.save(doc_path)
        logging.info(f"DOCX summary created: {doc_path}")
        return doc_path
    except Exception as e:
        logging.warning(f"DOCX creation failed: {e}")
    return None

def error_file(message):
    try:
        with open("/home/bots/modules/SC-ImmoCalcul/error.txt", "a", encoding="utf-8") as f:
            f.write(str(message) + "\n")
    except Exception:
        pass


async def create_context(pw, args, out_dir: Path):
    """Creates a robust browser context using Chromium browser on macOS."""
    viewport = {"width": args.viewport_width, "height": args.viewport_height}
    
    # Minimal launch args - avoid automation detection
    launch_args = [
        "--disable-blink-features=AutomationControlled",
        f"--window-size={args.viewport_width},{args.viewport_height}",
    ]
    
    # Use a dedicated automation profile
    automation_profile = os.path.expanduser("Chromium_Automation")
    
    logging.info(f"Using Chromium automation profile: {automation_profile}")
    logging.info("Using Playwright's bundled Chromium browser.")
    
    # Setup video recording directory if requested
    record_video_dir = None
    if getattr(args, 'record_video', False):
        try:
            vid_dir = out_dir / "video"
            vid_dir.mkdir(parents=True, exist_ok=True)
            record_video_dir = str(vid_dir)
            logging.info(f"Video recording enabled: {record_video_dir}")
        except Exception as e:
            logging.warning(f"Could not create video directory: {e}")
            record_video_dir = None
    
    try:
        # Load proxy from environment
        proxy_host = os.getenv("PROXY_HOST")
        proxy_port = os.getenv("PROXY_PORT")
        proxy_user = os.getenv("PROXY_USER")
        proxy_pass = os.getenv("PROXY_PASS")
        
        proxy_config = None
        if proxy_host and proxy_port:
            proxy_config = {
                "server": f"http://{proxy_host}:{proxy_port}",
                "username": proxy_user,
                "password": proxy_pass,
            }
            logging.info(f"Using proxy: {proxy_host}:{proxy_port}")
        
        context = await pw.chromium.launch_persistent_context(
            user_data_dir=automation_profile,
            headless=False,
            args=launch_args,
            viewport=viewport,
            java_script_enabled=True,
            locale="en-US",
            accept_downloads=True,
            proxy=proxy_config,
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/136.0.0.0 Safari/537.36",
            extra_http_headers={
                "sec-ch-ua": '"Chromium";v="136", "Not-A.Brand";v="24", "Google Chrome";v="136"',
                "sec-ch-ua-mobile": "?0",
                "sec-ch-ua-platform": '"Windows"',
            },
            **({"record_video_dir": record_video_dir} if record_video_dir else {})
        )
        # Minimal anti-bot detection script
        await context.add_init_script("""
            // Remove webdriver property
            Object.defineProperty(navigator, 'webdriver', {
                get: () => undefined
            });
            
            // Remove automation-related window properties
            delete window.cdc_adoQpoasnfa76pfcZLmcfl_Array;
            delete window.cdc_adoQpoasnfa76pfcZLmcfl_Promise;
            delete window.cdc_adoQpoasnfa76pfcZLmcfl_Symbol;
        """)
        
        logging.info("Browser context created successfully!")
        return context
        
    except Exception as e:
        logging.error(f"Failed to create browser context: {e}", exc_info=True)
        raise


async def robust_autocomplete(page: Page, search_field: Locator, query: str):
    """Fills search field and robustly clicks autocomplete suggestion."""
    await search_field.click()
    await search_field.fill("")
    await search_field.type(query, delay=70)

    try:
        await page.wait_for_load_state("networkidle", timeout=60000)
    except PWTimeoutError:
        logging.debug("Network did not go idle after typing, proceeding anyway.")

    suggestion = page.locator(AUTOCOMPLETE_ADDR_XPATH).first
    for attempt in range(7):
        if await suggestion.count() > 0 and await suggestion.is_visible():
            logging.info(f"Autocomplete suggestion found on attempt {attempt+1}. Clicking it.")
            await suggestion.click()
            return
        logging.debug(f"Suggestion not visible on attempt {attempt+1}. Nudging with ArrowDown.")
        await page.keyboard.press("ArrowDown")
        await asyncio.sleep(0.4)

    logging.warning("Autocomplete suggestion not found after multiple attempts. Pressing 'Enter' as fallback.")
    await search_field.press("Enter")


async def do_sequence(args) -> Dict[str, Any]:
    # Create unique output directory for run data
    run_id = args.run_id or datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    out_dir = Path("run_steps") / run_id
    out_dir.mkdir(parents=True, exist_ok=True)
    
    logging.info(f"Output directory: {out_dir}")

    email = args.email or os.getenv("IMMOCALCUL_EMAIL")
    password = args.password or os.getenv("IMMOCALCUL_PASSWORD")
    if not email or not password:
        set_step("invalid-credentials")
        logging.critical("Missing IMMOCALCUL_EMAIL or IMMOCALCUL_PASSWORD in environment.")
        raise ValueError("Missing email or password credentials.")

    # Setup virtual display if requested
    xvfb_process = None
    if getattr(args, 'virtual_display', False) and not sys.platform.startswith('darwin'):
        try:
            logging.info("Starting virtual display (Xvfb)...")
            xvfb_process = subprocess.Popen(
                ['Xvfb', ':99', '-screen', '0', '1366x768x24'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL
            )
            os.environ['DISPLAY'] = ':99'
            logging.info("Virtual display started on :99")
            await asyncio.sleep(1)  # Give Xvfb time to start
        except FileNotFoundError:
            logging.warning("Xvfb not found. Install with: apt-get install xvfb (Linux only)")
            xvfb_process = None
        except Exception as e:
            logging.warning(f"Could not start virtual display: {e}")
            xvfb_process = None
    elif getattr(args, 'virtual_display', False) and sys.platform.startswith('darwin'):
        logging.warning("Virtual display (Xvfb) is only supported on Linux. Skipping for macOS.")

    summary: Dict[str, Any] = {
        "email_used": email,
        "mode": "lot" if args.lot else "address",
        "lot": args.lot,
        "address_line_1": "",
        "address_line_2": "",
        "Other Street Number": "",
        "Other Street": "",
        "Other Unit": "",
        "Other City": "",
        "Other State": "Quebec",  # Hardcoded
        "Other Zip": "",
        "Other Country": "Canada",  # Hardcoded
        "classification": None,
        "char6": "",
        "amounts": [],
        "Price": None,
        "resume_text": "",
        "active_text_main": "",
        "active_text_measures": "",
        "active_text_avis": "",
        "risk_issues": [],
        "Analyse risque": "",
        "comparables_text": "",
        "main_photo_local_path": None,
        "Picture 1": "",
        "Google Drive": "",
        "map_layers": [],
        "uploaded_files": {}
    }

    async with async_playwright() as pw:
        context = await create_context(pw, args, out_dir)
        
        # Start tracing if requested
        if args.trace:
            await context.tracing.start(screenshots=True, snapshots=True)
            logging.info("Playwright tracing started")
        
        page = context.pages[0] if context.pages else await context.new_page()
        page.set_default_timeout(args.selector_timeout)
        page.set_default_navigation_timeout(args.navigation_timeout)

        # Add diagnostics for headless runs
        if args.headless:
            page.on("console", lambda m: logging.debug(f"[console:{m.type}] {m.text}"))
            page.on("response", lambda r: logging.debug(f"[resp] {r.status} {r.url}"))

        sc = StepCapture(out_dir)

        try:
            # --- Auth with Improved Login Check ---
            set_step("auth")
            await page.goto("https://immocalcul.com/immobilier", timeout=90000, wait_until="domcontentloaded")
            await asyncio.sleep(1.2) # Extra settle time for scripts

            # Check for the account/subscription info `div` to see if we're logged in.
            account_div_selector = "//div[contains(@class,'map_abonnement__')]"
            try:
                await page.locator(account_div_selector).first.wait_for(timeout=50000)
                logging.info("Account info div found. Already logged in.")
            except PWTimeoutError:
                logging.info("Account info div not found. Performing login...")
                await page.locator("//span[contains(text(),'Connexion')]").first.click()
                await page.fill("//input[@type='email']", email)
                await page.fill("//input[@type='password']", password)
                await page.locator("(//button[contains(text(),'Connexion')])[last()]").click()

            # Wait for the main page to be ready after potential login
            await page.wait_for_selector("#searchField, #searchFieldLot", timeout=90000)
            logging.info("Main page is ready.")

            # --- Robust Search with Human-like Typing ---
            set_step("search")
            if args.lot:
                search_selector = "#searchFieldLot"
                search_query = args.lot
            else:
                search_selector = "#searchField"
                search_query = f"{args.address_number or ''} {args.address_street or ''} {args.address_city or ''}".strip()

            logging.info(f"Searching for: '{search_query}' with robust autocomplete.")
            search_field = page.locator(search_selector)
            await robust_autocomplete(page, search_field, search_query)

            # # Wait for property content to load after search selection
            try:
            #     # Wait for autocomplete suggestion and click it if available
                suggestion = page.locator(".map_adresseproperty__Q7GLP").first
                await suggestion.wait_for(state="visible", timeout=8000)
                await suggestion.click()
                logging.info("Clicked address div.")
                # Wait for property panel to become naturally visible
                await asyncio.sleep(3)
                await page.wait_for_load_state("networkidle", timeout=30000)
            except Exception:
                logging.info("No autocomplete suggestion, trying Enter fallback.")
                search_field = page.locator("#searchFieldLot")
                await search_field.press("Enter")

            



                # Wait for panel to be visible
                panel_visible = False
                for selector in [
                    "[class*='dataAdress_modalPropertySession']",
                    "[class*='dataAdress_blocDataAdress']",
                    "[class*='dataAdress_adresseproperty']",
                ]:
                    try:
                        el = page.locator(selector).first
                        await el.wait_for(state="visible", timeout=60000)
                        panel_visible = True
                        logging.info(f"Property panel visible: {selector}")
                        break
                    except Exception:
                        continue

                if not panel_visible:
                    raise RuntimeError("Property panel not found in DOM.")

                await asyncio.sleep(1)
                logging.info("Property page content is loaded.")

            except RuntimeError:
                set_step("not-found")
                logging.error("Property found, but content (resume) did not load in time.")
                raise

            # --- Certificate (best-effort) ---
            set_step("certificate")
            logging.info("Attempting to download certificate.")
            try:
                # Wait for new page/tab to open when clicking certificate
                async with context.expect_page() as new_page_info:
                    await page.locator("(//span[contains(@class, 'dataAdress_iconTitleDataTwo')])[2]").click()
                
                certificate_page = await new_page_info.value
                await certificate_page.wait_for_load_state('domcontentloaded', timeout=args.navigation_timeout)
                cert_url = certificate_page.url
                logging.info(f"Certificate opened in new tab: {cert_url}")
                
                filename = cert_url.split('/')[-1]
                if not filename.endswith('.pdf'):
                    filename = (filename.split('=')[1] if '=' in filename else filename) + '.pdf'

                pdf_path_certificate = out_dir / filename
                response = requests.get(cert_url, timeout=30)
                response.raise_for_status()
                with open(pdf_path_certificate, 'wb') as f:
                    f.write(response.content)
                summary["certificate_local_path"] = str(pdf_path_certificate)
                logging.info(f"Certificate downloaded and saved: {pdf_path_certificate}")
                await certificate_page.close()
                
            except Exception as e:
                set_step("certificate-error")
                logging.warning(f"Certificate download failed: {e}")
                # Try to close any orphaned tabs
                for p in context.pages[1:]:
                    if p != page: await p.close()

            # --- Extract all other data as before ---
            set_step("complete-fields")
            lot_number_loc = page.locator("(//span[contains(@class, 'dataAdress_iconTitleDataTwo')])[1]").first
            address_h1_loc = page.locator("//h1[contains(@class,'dataAdress_adresseproperty__pJjrz')]").first
            address_line2_loc = address_h1_loc.locator("span")

            summary["lot"] = (await get_text_from_locator(lot_number_loc)) or summary["lot"]
            summary["lot"] = re.sub(r'\s+', '', summary["lot"] or "")
            h1_full_text = await get_text_from_locator(address_h1_loc)
            line2 = await get_text_from_locator(address_line2_loc)
            line1 = (h1_full_text or "").replace(line2 or "", "").strip()
            summary["address_line_1"] = line1
            summary["address_line_2"] = line2

            # (Address parsing logic remains the same)
            unit_match = re.search(r'(#\s?\w+)$', line1 or "")
            if unit_match:
                summary["Other Unit"] = unit_match.group(1).strip()
                line1 = line1[:unit_match.start()].strip()
            street_num_match = re.match(r'^(\d+(?:-\d+)?)', line1 or "")
            if street_num_match:
                summary["Other Street Number"] = street_num_match.group(1).strip()
                summary["Other Street"] = (line1 or "")[street_num_match.end():].strip()
            else:
                summary["Other Street"] = line1 or ""
            parts = [p.strip() for p in (line2 or "").split(',')]
            if len(parts) >= 3:
                summary["Other City"] = parts[0]
                summary["Other State"] = "Quebec"
                summary["Other Zip"] = parts[-1]
            else:
                summary["Other State"] = "Quebec"
            summary["Other Country"] = "Canada"

            summary["resume_text"] = await get_text_from_locator(page.locator(RESUME_XPATH).first)
            summary["char6"] = await get_text_from_locator(page.locator(CHAR6_XPATH).first)
            amounts = parse_amounts(summary["resume_text"]); summary["amounts"] = amounts
            summary["classification"] = classify(amounts, summary["char6"]); summary["Price"] = sum(amounts) if amounts else None
            logging.info(f"Property classified as '{summary['classification']}' with price {summary['Price']}")
            summary["active_text_main"] = await get_all_text_active(page)

            # --- Main tab capture ---
            await page.wait_for_timeout(200)
            await sc.shot_xpath_zoomfit(page, DATA_SECTION_XPATH, "main_tab_zoom33", zoom=0.50)
            photo_path = await capture_main_photo(page, out_dir, sc)
            summary["main_photo_local_path"] = photo_path

            # --- Measures tab ---
            try:
                # Try to click the tab by text (supports both English and French)
                measures_tab = page.locator("//span[text()='Measures' or text()='Mesures']").first
                await measures_tab.click()
                try:
                    await page.wait_for_load_state("networkidle", timeout=80000)
                except Exception:
                    logging.info("Network did not go idle after clicking Measures tab, proceeding anyway.")
                summary["active_text_measures"] = await get_all_text_active(page)
                await page.wait_for_timeout(200)
                await sc.shot_xpath_zoomfit(page, DATA_SECTION_XPATH, "measures_tab_zoom33", zoom=0.50)
            except Exception as e:
                logging.warning(f"Measures tab failed: {e}")

            # --- Avis tab ---
            set_step("avis")
            try:
                avis_tab = page.locator("//span[text()='Avis' or text()='Notices']").first
                await avis_tab.click()
                try:
                    await page.wait_for_load_state("networkidle", timeout=80000)
                except Exception:
                    logging.info("Network did not go idle after clicking Avis tab, proceeding anyway.")
                summary["active_text_avis"] = await get_all_text_active(page)
                risk_issues = await get_risk_issue_headings(page); summary["risk_issues"] = risk_issues
                summary["Analyse risque"] = ", ".join(risk_issues)
                await page.wait_for_timeout(200)
                await sc.shot_xpath_zoomfit(page, DATA_SECTION_XPATH, "avis_tab_zoom33", zoom=0.35)
            except Exception as e:
                set_step("avis-parse-error")
                logging.warning(f"Avis tab failed or notices unreadable: {e}")

            # --- Comparables dialog ---
            try:
                comp_trigger = page.locator(COMPARABLES_TRIGGER_XPATH).first
                if await comp_trigger.is_visible(timeout=60000):
                    await comp_trigger.click()
                    await page.wait_for_load_state("networkidle", timeout=60000)
                    await page.wait_for_selector(COMPARABLES_DIALOG_CONTAINER_XPATH, timeout=60000)
                    comp_dialog = page.locator(COMPARABLES_DIALOG_CONTENT_XPATH).first
                    summary["comparables_text"] = await get_text_from_locator(comp_dialog)
                else:
                    logging.info("Comparables trigger not visible.")
            except Exception as e:
                logging.warning(f"Comparables dialog capture failed: {e}")

            # Maps flow
            try:
                MAP_OPEN_XPATH = "//div//span[contains(., 'Voir sur la carte')]"
                await page.locator(MAP_OPEN_XPATH).click()
                logging.info("Waiting for map initialization...")
                await page.wait_for_load_state("networkidle")
                await asyncio.sleep(2) # Allow map tiles to render

                try:
                    ferme_button = page.get_by_label("Fermé")
                    for _ in range(2):
                        if await ferme_button.is_visible(timeout=2000):
                            await ferme_button.click()
                            await asyncio.sleep(0.5)
                except Exception:
                    pass

                await sc.shot_xpath_zoomfit(page, "//main[@id='main']/div[contains(@class,'map_blocvueSession__')]//div[contains(@class,'map_blocvueIntoSession__')]", "map_opened_zoom33", zoom=0.80)
                canvas_loc = page.locator(MAP_CANVAS_XPATH).last
                await canvas_loc.wait_for(timeout=10000)
                # Wait for canvas to be properly rendered
                await page.wait_for_function("() => document.querySelector('canvas')?.getContext('2d')")


                # Hide overlays before canvas capture
                overlay_xpaths = [
                    "//div[contains(@class,'map_cardResultMap')]",
                    "//div[contains(@class,'map_customMapControls')]",
                    "//div[contains(@class,'grecaptcha-logo')]/iframe",
                    "//div[contains(@class,'map_mapAlert')]/div",
                    "//div[contains(@class,'noprint')]/div[contains(@class,'map_mapcontrolOptionSession')]"
                ]
                await page.evaluate(f"xpaths => {{ for (const xpath of xpaths) {{ const result = document.evaluate(xpath, document, null, XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null); for (let i = 0; i < result.snapshotLength; i++) {{ result.snapshotItem(i).style.display = 'none'; }} }} }}", overlay_xpaths)

                img_path = out_dir / f"02_map_opened_zoom33.png"
                await canvas_loc.screenshot(path=str(img_path))

                # Restore overlays
                await page.evaluate(f"xpaths => {{ for (const xpath of xpaths) {{ const result = document.evaluate(xpath, document, null, XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null); for (let i = 0; i < result.snapshotLength; i++) {{ result.snapshotItem(i).style.display = ''; }} }} }}", overlay_xpaths)
                
                # Close map
                await page.locator("//div[contains(@class,'map_categoriesItemListing__')]/span").click()
            except Exception as e:
                logging.warning(f"Map flow failed: {e}")

            # --- Satellite toggle (best-effort) ---
            try:
                sat_btn = page.locator(SATELLITE_BTN_XPATH).first
                if await sat_btn.is_visible(timeout=4000):
                    await sat_btn.click()
                    await page.wait_for_load_state("networkidle")
            except Exception as e:
                logging.warning(f"Satellite toggle failed: {e}")

            # --- Map layers (best-effort; if map not visible, loop will quietly skip) ---
            for idx, label in enumerate(MAP_ITEMS, start=1):
                layer_loc = page.locator(MAP_LAYER_BASE_XPATH.format(idx)).first
                try:
                    if await layer_loc.is_visible(timeout=5000):
                        await layer_loc.click(force=True)
                        await page.wait_for_load_state("networkidle", timeout=8000)
                        await asyncio.sleep(1) # Final render wait

                        canvas_loc = page.locator(MAP_CANVAS_XPATH).last
                        await canvas_loc.wait_for(timeout=6000)

                        overlay_xpaths = [
                            "//div[contains(@class,'map_cardResultMap')]",
                            "//div[contains(@class,'map_customMapControls')]",
                            "//div[contains(@class,'grecaptcha-logo')]/iframe",
                            "//div[contains(@class,'map_mapAlert')]/div",
                            "//div[contains(@class,'noprint')]/div[contains(@class,'map_mapcontrolOptionSession')]"
                        ]
                        await page.evaluate(f"xpaths => {{ for (const xpath of xpaths) {{ const result = document.evaluate(xpath, document, null, XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null); for (let i = 0; i < result.snapshotLength; i++) {{ result.snapshotItem(i).style.display = 'none'; }} }} }}", overlay_xpaths)

                        img_path = out_dir / f"map_canvas_{idx:02d}_{label.replace(' ','_')}.png"
                        await canvas_loc.screenshot(path=str(img_path))

                        # Restore overlays
                        await page.evaluate(f"xpaths => {{ for (const xpath of xpaths) {{ const result = document.evaluate(xpath, document, null, XPathResult.ORDERED_NODE_SNAPSHOT_TYPE, null); for (let i = 0; i < result.snapshotLength; i++) {{ result.snapshotItem(i).style.display = ''; }} }} }}", overlay_xpaths)

                        overlay_label(img_path, label)
                        summary["map_layers"].append({"index": idx, "label": label, "file": str(img_path)})
                except Exception as e:
                    logging.error(f"Failed to capture layer {idx} ({label}): {e}")

            # --- Outputs ---
            pdf_path = write_pdf_from_screenshots(out_dir, summary)
            docx_path = create_docx_summary(out_dir, summary)

            # --- Google Drive upload (optional) ---
            set_step("drive-upload")
            if True:
                try:
                    service = get_drive_service()
                    if service:
                        address_name = summary['address_line_1']
                        folder_name = address_name if address_name else f"ImmoCalcul_{datetime.now().strftime('%Y%m%d')}"
                        subfolder_id = args.sub_folder_id #create_drive_folder(service, folder_name, PARENT_FOLDER_ID)
                        summary["Google Drive"] = get_folder_url(subfolder_id)
                        logging.info(f"Created Google Drive folder: {summary['Google Drive']}")

                        if summary["main_photo_local_path"]:
                            photo_id, photo_link = upload_file_to_drive(service, summary["main_photo_local_path"], subfolder_id)
                            if photo_id:
                                summary['uploaded_files']['main_photo'] = {'id': photo_id, 'link': photo_link}
                                summary['Picture 1'] = f"https://drive.google.com/uc?id={photo_id}"
                        if pdf_path:
                            pdf_id, pdf_link = upload_file_to_drive(service, pdf_path, subfolder_id)
                            if pdf_id:
                                summary['uploaded_files']['pdf_report'] = {'id': pdf_id, 'link': pdf_link}
                        if 'pdf_path_certificate' in locals() and pdf_path_certificate:
                            try:
                                cert_id, cert_link = upload_file_to_drive(service, pdf_path_certificate, subfolder_id)
                                if cert_id:
                                    summary['uploaded_files']['certificate'] = {'id': cert_id, 'link': cert_link}
                            except Exception:
                                pass
                        if docx_path:
                            docx_id, docx_link = upload_file_to_drive(service, docx_path, subfolder_id)
                            if docx_id:
                                summary['uploaded_files']['docx_summary'] = {'id': docx_id, 'link': docx_link}
                except Exception as e:
                    logging.warning(f"Google Drive integration failed: {e}", exc_info=True)
            else:
                logging.info("No PARENT_DRIVE_FOLDER_ID set. Skipping Google Drive upload.")

            summary_file_path = out_dir / "summary.json"
            summary_file_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")

            logging.info(f"Run finished. Summary file at: {summary_file_path}")
            print(summary_file_path, flush=True)
            return summary

        except Exception as e:
            set_step("error")
            logging.error(f"A fatal error occurred: {e}", exc_info=True)
            err_file = out_dir / "error.txt"
            err_file.write_text(traceback.format_exc(), encoding="utf-8")
            if args.debug_dump:
                try:
                    html_dump = out_dir / "last_page.html"
                    await page.screenshot(path=str(out_dir / "error_screenshot.png"))
                    html_dump.write_text(await page.content(), encoding="utf-8")
                except Exception: pass
            raise
        finally:
            if args.trace:
                try:
                    trace_path = out_dir / "trace.zip"
                    await context.tracing.stop(path=str(trace_path))
                    logging.info(f"Playwright trace saved: {trace_path}")
                except Exception: pass
            
            # Finalize video recording if enabled
            if getattr(args, 'record_video', False):
                try:
                    vid_dir = out_dir / "video"
                    for idx, p in enumerate(context.pages, start=1):
                        try:
                            if hasattr(p, 'video') and p.video:
                                try:
                                    vpath = await p.video.path()
                                except Exception:
                                    vpath = None
                                if vpath:
                                    src = Path(vpath)
                                    dest = vid_dir / f"page_{idx:02d}.webm"
                                    try:
                                        shutil.move(str(src), str(dest))
                                        logging.info(f"Saved video for page {idx}: {dest}")
                                    except Exception as e:
                                        logging.warning(f"Could not move video {src} to {dest}: {e}")
                        except Exception:
                            continue
                except Exception as e:
                    logging.warning(f"Finalizing recorded videos failed: {e}")
            
            await context.close()
    
    # Cleanup virtual display if it was started
    if xvfb_process:
        try:
            xvfb_process.terminate()
            xvfb_process.wait(timeout=5)
            logging.info("Virtual display stopped.")
        except Exception as e:
            logging.warning(f"Could not stop virtual display: {e}")
            try:
                xvfb_process.kill()
            except Exception:
                pass

def build_arg_parser():
    p = argparse.ArgumentParser(description="Full step-by-step ImmoCalcul scraper using personal Google Drive.")
    group = p.add_mutually_exclusive_group(required=True)
    group.add_argument("--lot", help="Cadastral lot number")
    group.add_argument("--address-number", help="Street number (with --address-street and --address-city)")
    p.add_argument("--address-street")
    p.add_argument("--address-city")
    p.add_argument("--email", help="Login email (overrides IMMOCALCUL_EMAIL)")
    p.add_argument("--password", help="Login password (overrides IMMOCALCUL_PASSWORD)")
    p.add_argument("--parent-folder-id", help="Google Drive parent folder ID (overrides PARENT_DRIVE_FOLDER_ID env var)")
    p.add_argument("--sub-folder-id", help="Google Drive sub folder ID (use existing folder, don't create new)")
    p.add_argument("--run-id", help="Unique run ID for output directory (default: auto-generated)")
    p.add_argument("--delay-min", type=float, default=0.5, help="Minimum random delay (seconds)")
    p.add_argument("--delay-max", type=float, default=1.5, help="Maximum random delay (seconds)")
    p.add_argument("--fixed-delay", type=float, help="Fixed delay (overrides random range)")
    p.add_argument("--viewport-width", type=int, default=1366, help="Viewport width for headless")
    p.add_argument("--viewport-height", type=int, default=768, help="Viewport height for headless")
    p.add_argument("--headless", action="store_true", default=False, help="Run headless (default: headed).")
    p.add_argument("--virtual-display", action="store_true", default=False, help="Use virtual display (Xvfb) for Linux VPS simulation (headed mode with virtual screen).")
    p.add_argument("--trace", action="store_true", default=False, help="Record Playwright trace (trace.zip) for debugging")
    p.add_argument("--record-video", action="store_true", default=False, help="Record a screen video for the run (saved under run_steps/<run_id>/video)")
    p.add_argument("--selector-timeout", type=int, default=60000, help="Selector timeout ms")
    p.add_argument("--navigation-timeout", type=int, default=90000, help="Navigation timeout ms")
    return p


def main():
    global PARENT_FOLDER_ID
    setup_logging()
    args = build_arg_parser().parse_args()

    if args.parent_folder_id:
        PARENT_FOLDER_ID = args.parent_folder_id

    if not PARENT_FOLDER_ID:
        PARENT_FOLDER_ID = "1AubnRZyvIBTiWwjiG0UAXsHyKE4wqw47"  # Default to a known folder for testing, but should be overridden in production
        logging.warning(f"PARENT_DRIVE_FOLDER_ID was not set. now set to {PARENT_FOLDER_ID}.")
        

    if args.fixed_delay is not None and args.fixed_delay < 0:
        sys.exit("Fixed delay must be >= 0")
    if args.delay_min > args.delay_max:
        sys.exit("delay-min cannot exceed delay-max")
    if args.address_number and (not args.address_street or not args.address_city):
        sys.exit("Must supply --address-street and --address-city with --address-number")

    try:
        asyncio.run(do_sequence(args))
    except KeyboardInterrupt:
        logging.warning("Interrupted by user.")
        sys.exit(130)
    except Exception:
        sys.exit(1)


if __name__ == "__main__":
    main()

